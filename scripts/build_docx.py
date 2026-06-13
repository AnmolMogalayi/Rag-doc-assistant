"""Build a single Word (.docx) document from the project's Markdown docs.

Handles headings, paragraphs, blockquotes, bullet/numbered/checkbox lists, tables, fenced
code blocks, horizontal rules, and inline **bold** / `code`. Not a full CommonMark parser —
just enough to render this repo's docs cleanly.

Usage:
    python -m scripts.build_docx                      # default bundle -> docs/RAG_Documentation.docx
    python -m scripts.build_docx --only HOW_TO_RUN    # single file
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

# Files combined into the default bundle, in order.
DEFAULT_BUNDLE = [
    ("HOW_TO_RUN.md", "How to Run — Setup Guide & Manual Steps"),
    ("PHASE1_ANALYSIS.md", "Phase 1 — Analysis & Architecture"),
    ("PHASE3_TESTING_AND_TRACEABILITY.md", "Phase 3 — Testing & Traceability"),
]

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.*)$")
NUMBERED_RE = re.compile(r"^(\s*)\d+\.\s+(.*)$")
CHECKBOX_RE = re.compile(r"^\s*[-*]\s+\[([ xX])\]\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def _shade(cell, hex_color: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc.append(shd)


def add_inline(paragraph, text: str) -> None:
    """Render **bold** and `code` inline within a paragraph."""
    # Split on bold and inline-code while keeping delimiters.
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            paragraph.add_run(tok)


def add_code_block(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # light gray shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            text = row[c_idx] if c_idx < len(row) else ""
            add_inline(para, text)
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _shade(cell, "2E6DA4")


def parse_table_block(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows, i = [], start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        if TABLE_SEP_RE.match(lines[i]):
            i += 1
            continue
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def convert_markdown(doc, md: str) -> None:
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            j = i + 1
            buf = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            add_code_block(doc, buf)
            i = j + 1
            continue

        # Table
        if line.lstrip().startswith("|") and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            rows, ni = parse_table_block(lines, i)
            add_table(doc, rows)
            doc.add_paragraph()
            i = ni
            continue

        # Horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            doc.add_paragraph().add_run("").add_break()
            i += 1
            continue

        # Heading
        m = HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            h = doc.add_heading(level=level)
            add_inline(h, text)
            i += 1
            continue

        # Blockquote
        if line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, text)
            i += 1
            continue

        # Checkbox list
        cm = CHECKBOX_RE.match(line)
        if cm:
            checked = cm.group(1).lower() == "x"
            p = doc.add_paragraph(style="List Bullet")
            box = p.add_run(("☑ " if checked else "☐ "))
            box.font.name = "Segoe UI Symbol"
            add_inline(p, cm.group(2))
            i += 1
            continue

        # Bullet list (with simple nesting via indentation)
        bm = BULLET_RE.match(line)
        if bm:
            indent = len(bm.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, bm.group(2))
            i += 1
            continue

        # Numbered list
        nm = NUMBERED_RE.match(line)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, nm.group(2))
            i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_inline(p, line.strip())
        i += 1


def build(bundle: list[tuple[str, str]], out_path: Path) -> None:
    doc = Document()

    # Base styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAG-Based Technical Documentation Assistant")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Self-corrective RAG • LangGraph + FastAPI\n"
                    "Express Analytics — AI/ML Engineer Intern Take-Home")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    toc = doc.add_paragraph()
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc.add_run("Contents: " + "  |  ".join(t for _, t in bundle)).font.size = Pt(9)

    for idx, (fname, _title) in enumerate(bundle):
        path = DOCS / fname
        if not path.exists():
            print(f"  ! skip missing {path}")
            continue
        doc.add_page_break()
        convert_markdown(doc, path.read_text(encoding="utf-8"))
        print(f"  + added {fname}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\nSaved: {out_path}")


def main() -> int:
    ap = argparse.ArgumentParser(description="Build a .docx from project Markdown docs.")
    ap.add_argument("--only", help="Base name (without .md) of a single docs/ file to convert")
    ap.add_argument("--out", help="Output .docx path")
    args = ap.parse_args()

    if args.only:
        name = args.only if args.only.endswith(".md") else f"{args.only}.md"
        bundle = [(name, name.replace(".md", ""))]
        out = Path(args.out) if args.out else DOCS / f"{name.replace('.md', '')}.docx"
    else:
        bundle = DEFAULT_BUNDLE
        out = Path(args.out) if args.out else DOCS / "RAG_Documentation.docx"

    print(f"Building {out.name} ...")
    build(bundle, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
